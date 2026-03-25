import sys
import os
import subprocess

def install_and_import(package, import_name):
    try:
        __import__(import_name)
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

install_and_import('python-docx', 'docx')
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

project_dir = r"c:\Users\vello\OneDrive\Desktop\imp\SACAIM\SACAIM25_Project"
doc_dir = os.path.join(project_dir, "Documentation")
doc_path = os.path.join(doc_dir, "Domain_Knowledge.docx")
out_path = os.path.join(doc_dir, "Domain_Knowledge_V3.docx")

doc = docx.Document(doc_path)

print("Cleaning old chapters...")
found = False
for p in list(doc.paragraphs):
    text_up = p.text.upper()
    if "CHAPTER 4" in text_up or "MODEL SELECTION, IMPLEMENTATION" in text_up:
        found = True
    if found:
        p_element = p._element
        p_element.getparent().remove(p_element)

doc.add_page_break()

def append_draft(draft_file):
    with open(os.path.join(doc_dir, draft_file), 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith("CHAPTER"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("[IMAGE:"):
            img_name = line.replace("[IMAGE:", "").replace("]", "").strip()
            img_path = os.path.join(project_dir, img_name.replace("/", "\\"))
            
            try:
                if os.path.exists(img_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(6.0))
                else:
                    p = doc.add_paragraph(f"[Missing Image: {img_name}]")
                    p.runs[0].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
            except Exception as e:
                p = doc.add_paragraph(f"[Error loading image {img_name}: {e}]")
                p.runs[0].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
        elif line.startswith("Fig"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.italic = True
            run.font.size = Pt(10)
        else:
            # Check if it starts with number dot (e.g., 4.1, 4.2)
            if len(line.split(" ")) > 1 and line.split(" ")[0].count(".") >= 1 and line.split(" ")[0][0].isdigit():
                heading = doc.add_heading(level=2)
                run = heading.add_run(line)
                run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            else:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

print("Appending Chapter 4...")
append_draft("ch4_draft.txt")

doc.add_page_break()

print("Appending Chapter 5...")
append_draft("ch5_draft.txt")
        
doc.save(out_path)
print("Finished compiling exhaustive documentation!")
