import sys
import os
import subprocess

def install_and_import(package, import_name):
    try:
        __import__(import_name)
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

install_and_import('python-docx', 'docx')
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

project_dir = r"c:\Users\vello\OneDrive\Desktop\imp\SACAIM\SACAIM25_Project"
doc_dir = os.path.join(project_dir, "Documentation")

doc_path = os.path.join(doc_dir, "Domain_Knowledge.docx")
draft_path = os.path.join(doc_dir, "draft_chapters.txt")
out_path = os.path.join(doc_dir, "Domain_Knowledge.docx")

print("Loading document...")
doc = docx.Document(doc_path)

# Insert a page break before starting Chapter 4
doc.add_page_break()

print("Reading draft...")
with open(draft_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line:
        continue
        
    if line.startswith("CHAPTER"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith("[IMAGE:"):
        img_name = line.replace("[IMAGE:", "").replace("]", "").strip()
        img_path = os.path.join(project_dir, img_name.replace("/", "\\"))
        
        try:
            if os.path.exists(img_path):
                # Add centered image
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(6.0))
            else:
                p = doc.add_paragraph(f"[Missing Image: {img_name}]")
                p.runs[0].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
        except Exception as e:
            p = doc.add_paragraph(f"[Error loading image {img_name}: {e}]")
            p.runs[0].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
    elif line.startswith("Fig:"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.italic = True
        run.font.size = Pt(10)
    else:
        # Check if heading
        headings = [
            "Introduction", "Model Selection and Implementation", 
            "Performance Evaluation", "Visual Analysis of Forecasts", 
            "Dashboard and API Integration", "Summary", 
            "Key Findings", "Limitations", "Future Scope"
        ]
        
        if line in headings:
            heading = doc.add_heading(level=2)
            run = heading.add_run(line)
            run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
        elif line == "MODEL SELECTION, IMPLEMENTATION AND RESULTS" or line == "CONCLUSION":
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
print("Saving document...")
doc.save(out_path)
print("Docx generation successful!")
